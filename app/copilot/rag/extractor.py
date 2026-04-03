"""
Text Extraction from Documents
Supports PDF, DOCX, TXT, MD, CSV
"""
import io
from typing import Optional


async def extract_text(file_content: bytes, file_type: str) -> str:
    """
    Extract text content from various document formats.

    Args:
        file_content: Raw file bytes
        file_type: File extension (pdf, docx, txt, md, csv, png, jpg, gif, webp)

    Returns:
        Extracted text content

    Raises:
        ValueError: If file type is not supported
    """
    extractors = {
        "pdf": extract_pdf,
        "docx": extract_docx,
        "txt": extract_txt,
        "md": extract_txt,  # Markdown is plain text
        "csv": extract_csv,
        "xls": extract_excel,
        "xlsx": extract_excel,
        # Image types - return placeholder for now (will use vision AI in future)
        "png": extract_image,
        "jpg": extract_image,
        "jpeg": extract_image,
        "gif": extract_image,
        "webp": extract_image,
    }

    extractor = extractors.get(file_type.lower())
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")

    return await extractor(file_content)


async def extract_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF files using PyMuPDF (fitz).
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF is required for PDF extraction. Install with: pip install pymupdf")

    text_parts = []

    # Open PDF from bytes
    doc = fitz.open(stream=file_content, filetype="pdf")

    for page_num, page in enumerate(doc):
        # Extract text from page
        text = page.get_text("text")
        if text.strip():
            text_parts.append(f"[Page {page_num + 1}]\n{text}")

    doc.close()

    return "\n\n".join(text_parts)


async def extract_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX files using python-docx.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX extraction. Install with: pip install python-docx")

    doc = Document(io.BytesIO(file_content))

    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n\n".join(text_parts)


async def extract_txt(file_content: bytes) -> str:
    """
    Extract text from plain text files.
    Handles various encodings.
    """
    # Try common encodings
    encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]

    for encoding in encodings:
        try:
            return file_content.decode(encoding)
        except UnicodeDecodeError:
            continue

    # Fallback: decode with errors ignored
    return file_content.decode("utf-8", errors="ignore")


async def extract_csv(file_content: bytes) -> str:
    """
    Extract text from CSV files.
    Converts to readable tabular format.
    """
    import csv

    text = await extract_txt(file_content)
    reader = csv.reader(io.StringIO(text))

    rows = []
    for row in reader:
        if any(cell.strip() for cell in row):
            rows.append(" | ".join(cell.strip() for cell in row))

    return "\n".join(rows)


async def extract_excel(file_content: bytes) -> str:
    """
    Extract text from Excel files (XLS, XLSX).
    Converts all sheets to readable tabular format.
    """
    try:
        import pandas as pd
    except ImportError:
        raise ImportError("pandas is required for Excel extraction. Install with: pip install pandas openpyxl")

    # Read Excel file from bytes
    excel_file = io.BytesIO(file_content)

    # Read all sheets
    excel_data = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')

    text_parts = []

    for sheet_name, df in excel_data.items():
        # Add sheet header
        text_parts.append(f"[Sheet: {sheet_name}]")

        # Convert dataframe to string with proper formatting
        # Replace NaN with empty string
        df = df.fillna('')

        # Convert to string representation with columns
        if not df.empty:
            # Get column headers
            headers = " | ".join(str(col) for col in df.columns)
            text_parts.append(headers)
            text_parts.append("-" * len(headers))

            # Add rows
            for _, row in df.iterrows():
                row_text = " | ".join(str(cell) for cell in row.values)
                if row_text.strip():
                    text_parts.append(row_text)

        text_parts.append("")  # Empty line between sheets

    return "\n".join(text_parts)


async def extract_image(file_content: bytes) -> str:
    """
    Extract information from image files.

    For now, returns a placeholder indicating an image is present.
    TODO: Integrate vision AI (GPT-4 Vision, Claude Vision, etc.) for actual image analysis.
    """
    import base64

    # Get image dimensions and format info
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(file_content))
        width, height = img.size
        format_name = img.format or "unknown"

        # For future vision AI integration, we can encode the image
        # base64_image = base64.b64encode(file_content).decode('utf-8')

        return f"[Image File: {format_name} format, {width}x{height} pixels]\n\nNote: This is an image file. To analyze its contents, a vision-capable AI model is required."
    except Exception:
        # Fallback if PIL is not available or image is corrupted
        return "[Image File]\n\nNote: This is an image file. To analyze its contents, a vision-capable AI model is required."


async def extract_with_metadata(file_content: bytes, file_type: str) -> dict:
    """
    Extract text with additional metadata (pages, sections, etc.).

    Returns:
        Dict with 'text', 'pages', 'metadata'
    """
    if file_type == "pdf":
        return await extract_pdf_with_metadata(file_content)
    else:
        text = await extract_text(file_content, file_type)
        return {
            "text": text,
            "pages": [],
            "metadata": {"file_type": file_type},
        }


async def extract_pdf_with_metadata(file_content: bytes) -> dict:
    """
    Extract text from PDF with page-level metadata.
    """
    try:
        import fitz
    except ImportError:
        raise ImportError("PyMuPDF is required for PDF extraction")

    doc = fitz.open(stream=file_content, filetype="pdf")

    pages = []
    full_text_parts = []

    for page_num, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            pages.append({
                "page_number": page_num + 1,
                "content": text,
                "char_start": len("\n\n".join(full_text_parts)) if full_text_parts else 0,
            })
            full_text_parts.append(text)

    # Get document metadata
    metadata = {
        "file_type": "pdf",
        "page_count": len(doc),
        "title": doc.metadata.get("title", ""),
        "author": doc.metadata.get("author", ""),
        "subject": doc.metadata.get("subject", ""),
        "creator": doc.metadata.get("creator", ""),
    }

    doc.close()

    return {
        "text": "\n\n".join(full_text_parts),
        "pages": pages,
        "metadata": metadata,
    }
